"""Build the fixed .docx fixtures used by the eval harness.

Run: `python evals/build_fixtures.py`
Also invoked automatically by `evals/run_eval.py` when fixtures are missing.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
FIXTURES = ROOT / "fixtures"


def _write(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for heading, paras in sections:
        doc.add_heading(heading, level=1)
        for p in paras:
            doc.add_paragraph(p)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"Wrote {path}")


def build_all() -> list[Path]:
    weak = FIXTURES / "weak_graphene.docx"
    _write(
        weak,
        "A Study of Graphene Based Sensors for Detecting Proteins",
        [
            (
                "Abstract",
                [
                    "In recent years, graphene has become very popular for a lot of sensing "
                    "applications, and in this paper we made a graphene sensor that can detect "
                    "proteins and it works really well and we believe it is significant for "
                    "diagnostics because it is obviously better than other methods, and it has "
                    "a lot of advantages that we will describe in order to show the performance "
                    "which we measured carefully in the lab."
                ],
            ),
            (
                "Introduction",
                [
                    "It is well known that detecting biomarkers is important. Graphene is a good "
                    "material. We used it to make a sensor. There are many papers about graphene "
                    "sensors, etc.",
                    "The sensitivity was very high and clearly the best, and we think this proves "
                    "our device is significant compared to previous works.",
                ],
            ),
            (
                "Methods",
                [
                    "We grew graphene and transferred it and then we functionalized it with a "
                    "linker molecule and attached antibodies to it in order to capture the "
                    "proteins , and we measured the electrical signal."
                ],
            ),
            (
                "Results",
                [
                    "The Dirac point moved a lot when we added the protein. The response was "
                    "very good and increased significantly with concentration."
                ],
            ),
            (
                "Conclusion",
                [
                    "In conclusion, we made a very good sensor that works really well and it is "
                    "obviously useful for a lot of applications."
                ],
            ),
        ],
    )

    claims = FIXTURES / "claims_heavy.docx"
    _write(
        claims,
        "Notes on Graphene Photodetector Performance",
        [
            (
                "Introduction",
                [
                    "It is well known that graphene photodetectors are obviously superior. "
                    "We believe our device is significant because the responsivity was clearly "
                    "the best ever reported, and we think this proves the mechanism without "
                    "needing further controls."
                ],
            ),
            (
                "Results",
                [
                    "The signal increased significantly. We believe this is significant for "
                    "on-chip sensing and obviously better than silicon."
                ],
            ),
        ],
    )
    return [weak, claims]


if __name__ == "__main__":
    build_all()
