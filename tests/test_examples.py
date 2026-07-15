"""Horizon 1 — before/after learning, style-card folding, and content-safety guard.

Uses MODEL_BACKEND=mock throughout so it needs no network/API keys/local LLM, but still
exercises the real diffing, retrieval-index, and severity-guard logic (only the final text
generation is mocked out elsewhere in the codebase, not any of this).
"""

from __future__ import annotations

import os

os.environ["MODEL_BACKEND"] = "mock"

from pathlib import Path

import pytest
from docx import Document

from professor_assistant import examples as examples_lib
from professor_assistant.config import get_settings
from professor_assistant.review import _enforce_content_flag_only, heuristic_review
from professor_assistant.docx_io import read_sections


def _write_docx(path: Path, heading: str, paragraphs: list[str]) -> None:
    d = Document()
    d.add_heading(heading, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(path))


@pytest.fixture()
def example_pair(tmp_path: Path) -> tuple[str, Path, Path]:
    before = tmp_path / "demo.before.docx"
    after = tmp_path / "demo.after.docx"
    _write_docx(
        before,
        "Introduction",
        [
            "It is well known that graphene is very popular for a lot of sensing "
            "applications, and we think this proves it is significant.",
        ],
    )
    _write_docx(
        after,
        "Introduction",
        [
            "Graphene field-effect transistors have emerged as a widely used platform "
            "for sensing applications, and these results suggest the approach is promising.",
        ],
    )
    return "demo", before, after


def test_find_example_pairs(tmp_path: Path, example_pair):
    pairs = examples_lib.find_example_pairs(tmp_path)
    assert len(pairs) == 1
    assert pairs[0][0] == "demo"


def test_find_example_pairs_ignores_unmatched(tmp_path: Path, example_pair):
    # An orphaned .before.docx with no matching .after.docx must not be picked up.
    (tmp_path / "orphan.before.docx").write_bytes(b"not a real docx, just needs to exist")
    pairs = examples_lib.find_example_pairs(tmp_path)
    assert [p[0] for p in pairs] == ["demo"]


def test_extract_patterns_from_pair_finds_edits(example_pair):
    pair, before, after = example_pair
    patterns = examples_lib.extract_patterns_from_pair(pair, before, after)
    assert patterns, "expected at least one edit pattern to be extracted"
    assert all(p.pair == pair for p in patterns)
    assert all(p.section == "introduction" for p in patterns)
    kinds = {p.kind for p in patterns}
    # The fixture deliberately contains a hedge-softening edit ("we think ... proves" ->
    # "these results suggest") and at least one plain phrase swap.
    assert "hedge_softened" in kinds


def test_summarize_and_format_examples_markdown(example_pair):
    pair, before, after = example_pair
    patterns = examples_lib.extract_patterns_from_pair(pair, before, after)
    summary = examples_lib.summarize_patterns(patterns)
    assert summary["patterns"] == len(patterns)
    assert summary["pairs"] == [pair]
    md = examples_lib.format_examples_markdown(summary)
    assert "Edit habits" in md


def test_summarize_empty_patterns_is_empty_markdown():
    assert examples_lib.format_examples_markdown(examples_lib.summarize_patterns([])) == ""


# --------------------------------------------------------------------------- #
# Must-fix 1: diff quality gate — regression tests against the known-garbage
# alignments produced by scripts/make_sample_examples.py before the fix.
# --------------------------------------------------------------------------- #


@pytest.fixture()
def garbage_pair(tmp_path: Path) -> tuple[str, Path, Path]:
    """Mirrors the real garbage found in the synthetic sample pairs: a fully
    restructured/reordered paragraph where word-level diffing coincidentally anchors on a
    shared word (e.g. "sensor") and hands back a nonsensical before/after "swap" spanning
    unrelated content, plus a run-on before-phrase that dangles across a clause boundary.
    """
    before = tmp_path / "garbage.before.docx"
    after = tmp_path / "garbage.after.docx"
    _write_docx(
        before,
        "Conclusion",
        [
            "In conclusion, we made a very good sensor that works really well and it is "
            "obviously useful for a lot of applications.",
        ],
    )
    _write_docx(
        after,
        "Conclusion",
        [
            "In conclusion, we demonstrated a GFET-based, label-free protein sensor that "
            "may be applicable to point-of-care diagnostics.",
        ],
    )
    return "garbage", before, after


@pytest.fixture()
def dangling_clause_pair(tmp_path: Path) -> tuple[str, Path, Path]:
    before = tmp_path / "dangling.before.docx"
    after = tmp_path / "dangling.after.docx"
    _write_docx(
        before,
        "Methods",
        [
            "We grew graphene and transferred it and then we functionalized it with a "
            "linker molecule and attached antibodies to it in order to capture the "
            "proteins, and we measured the electrical signal.",
        ],
    )
    _write_docx(
        after,
        "Methods",
        [
            "Graphene was grown by CVD, transferred onto the substrate, functionalized "
            "with a linker molecule, and conjugated with capture antibodies to bind the "
            "target protein; the electrical signal was then recorded.",
        ],
    )
    return "dangling", before, after


def test_diff_rejects_nonsensical_phrase_swap(garbage_pair):
    """'very good' -> 'GFET-based, label-free protein' must not survive as a phrase_swap:
    it pairs an unrelated short anchor with a comma-separated multi-descriptor blob, which
    would be applied literally (and nonsensically) by _example_swap_patterns() in review.py.
    """
    pair, before, after = garbage_pair
    patterns = examples_lib.extract_patterns_from_pair(pair, before, after)
    swaps = {(p.before.lower(), p.after.lower()) for p in patterns if p.kind == "phrase_swap"}
    assert ("very good", "gfet-based, label-free protein") not in swaps
    assert not any("," in b or "," in a for b, a in swaps), swaps


def test_diff_rejects_dangling_clause_swap(dangling_clause_pair):
    """'it and then we' -> 'onto the substrate' must not survive as a phrase_swap: the
    before-anchor dangles across a subject/clause boundary (contains 'and'/'then'), so it
    reads as nonsense on its own.
    """
    pair, before, after = dangling_clause_pair
    patterns = examples_lib.extract_patterns_from_pair(pair, before, after)
    swaps = {(p.before.lower(), p.after.lower()) for p in patterns if p.kind == "phrase_swap"}
    assert ("it and then we", "onto the substrate") not in swaps
    for before_phrase, _ in swaps:
        words = before_phrase.split()
        assert "and" not in words and "then" not in words, before_phrase


def test_diff_keeps_clean_phrase_swaps(example_pair):
    """A short, single-clause, non-listy swap (the fixture's real edit habit) must still
    survive the quality gate — the gate should reject garbage, not everything.
    """
    pair, before, after = example_pair
    patterns = examples_lib.extract_patterns_from_pair(pair, before, after)
    kinds = {p.kind for p in patterns}
    assert "hedge_softened" in kinds or "phrase_swap" in kinds


def test_classify_replace_accepts_clean_swap():
    assert examples_lib._classify_replace("very popular", "a widely used platform") == "phrase_swap"
    assert examples_lib._classify_replace("made", "demonstrated") == "phrase_swap"


def test_classify_replace_rejects_multi_clause_and_listy_phrases():
    # Comma-separated multi-descriptor blob paired with an unrelated short anchor.
    assert examples_lib._classify_replace("very good", "GFET-based, label-free protein") != "phrase_swap"
    # Anchor spans a clause boundary via a coordinating conjunction.
    assert examples_lib._classify_replace("it and then we", "onto the substrate") != "phrase_swap"
    assert examples_lib._classify_replace("we grew graphene and", "graphene was grown by CVD") != "phrase_swap"
    # Comma inside an otherwise short "list" of before/after clauses.
    assert (
        examples_lib._classify_replace("proteins, and we measured", "target protein")
        != "phrase_swap"
    )


def test_sample_pairs_produce_no_garbage_swaps():
    """End-to-end regression: the real scripts/make_sample_examples.py pairs (if present in
    data/examples/) must not surface any of the previously-shipped garbage swaps.
    """
    known_garbage = {
        ("very good", "gfet-based, label-free protein"),
        ("it and then we", "onto the substrate"),
        ("we grew graphene and", "graphene was grown by cvd"),
        ("proteins, and we measured", "target protein"),
        ("right one for future work, etc", "device under waveguide-coupled illumination"),
        ("in recent years, graphene has", "graphene field-effect transistors (gfets) have"),
    }
    patterns = examples_lib.extract_all_patterns()
    if not patterns:
        pytest.skip("no data/examples/ pairs present (run scripts/make_sample_examples.py)")
    swaps = {(p.before.lower(), p.after.lower()) for p in patterns if p.kind == "phrase_swap"}
    assert not (swaps & known_garbage), swaps & known_garbage
    assert not any("," in b or "," in a for b, a in swaps), swaps


# --------------------------------------------------------------------------- #
# Must-fix 3: style card / prompt honesty about synthetic vs. real examples.
# --------------------------------------------------------------------------- #


def test_examples_are_synthetic_detects_marker(tmp_path: Path, example_pair):
    assert examples_lib.examples_are_synthetic(tmp_path) is False
    (tmp_path / examples_lib.SYNTHETIC_MARKER_NAME).write_text("synthetic", encoding="utf-8")
    assert examples_lib.examples_are_synthetic(tmp_path) is True


def test_format_examples_markdown_labels_synthetic_data(example_pair):
    pair, before, after = example_pair
    patterns = examples_lib.extract_patterns_from_pair(pair, before, after)
    summary = examples_lib.summarize_patterns(patterns)

    honest_md = examples_lib.format_examples_markdown(summary, synthetic=True)
    assert "SYNTHETIC" in honest_md or "synthetic" in honest_md
    assert "these are his real editing decisions" not in honest_md.lower()
    # Must be explicit that it's NOT the real thing, not just silent about it.
    assert "not the professor's real edits" in honest_md.lower()

    real_md = examples_lib.format_examples_markdown(summary, synthetic=False)
    assert "SYNTHETIC" not in real_md
    # Even for real pairs, the card must not overclaim beyond "his editing decisions".
    assert "his real editing decisions" not in real_md


def test_make_sample_examples_writes_synthetic_marker(tmp_path: Path, monkeypatch):
    import importlib

    import scripts.make_sample_examples as make_sample_examples

    monkeypatch.setattr(make_sample_examples, "OUT_DIR", tmp_path)
    make_sample_examples.main()
    assert (tmp_path / examples_lib.SYNTHETIC_MARKER_NAME).exists()
    importlib.reload(make_sample_examples)  # restore module-level OUT_DIR for other tests


def test_build_and_retrieve_examples_index(tmp_path: Path, example_pair, monkeypatch):
    pair, before, after = example_pair
    settings = get_settings()
    monkeypatch.setattr(settings, "examples_dir", tmp_path)
    monkeypatch.setattr(settings, "examples_collection_name", f"examples_test_{pair}")

    result = examples_lib.build_examples_index(rebuild=True)
    assert result["pairs"] == 1
    assert result["patterns"] > 0
    assert result["collection_count"] == result["patterns"]

    matches = examples_lib.retrieve_examples(
        "graphene is very popular for a lot of sensing applications"
    )
    assert matches, "expected at least one retrieved edit pattern"
    assert matches[0].pair == pair

    context = examples_lib.format_examples_context(matches, prefix="E")
    assert context.startswith("[E1]")


def test_heuristic_review_content_is_flag_only():
    """The mock heuristic reviewer must never attach a rewrite to a 'content' suggestion."""
    _, sections = read_sections_from_text(
        "We think this proves the sensor is significant for diagnostics."
    )
    for section in sections:
        suggestions = heuristic_review(section)
        for s in suggestions:
            if s["severity"] == "content":
                assert s["suggestion"] == ""


def test_enforce_content_flag_only_strips_rewrites():
    """Defense-in-depth guard used for both mock and LLM paths."""
    suggestions = [
        {"severity": "content", "suggestion": "a sneaky silent rewrite", "applicable": True},
        {"severity": "language", "suggestion": "fine to keep", "applicable": True},
    ]
    _enforce_content_flag_only(suggestions)
    assert suggestions[0]["suggestion"] == ""
    assert suggestions[0]["applicable"] is False
    assert suggestions[1]["suggestion"] == "fine to keep"
    assert all("grounded_in" in s for s in suggestions)


def read_sections_from_text(text: str, tmp_path: Path | None = None):
    """Small helper: build a throwaway docx and read it back into Sections."""
    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = Path(d) / "draft.docx"
        _write_docx(path, "Introduction", [text])
        return read_sections(path)
