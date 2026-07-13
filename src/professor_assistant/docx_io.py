"""Read a student draft (.docx) into sections, and write v1 review outputs:
- reviewed.docx : original text with inline review notes (reason + suggestion)
- suggestions.md: a section-by-section redline

Note: v1 uses clearly-labelled inline note paragraphs (reliable across Word versions).
Native tracked-changes / margin comments are a planned v2 upgrade.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

# Canonical section types and the headings that map to them.
SECTION_KEYWORDS = {
    "abstract": "abstract",
    "introduction": "introduction",
    "background": "introduction",
    "methods": "methods",
    "method": "methods",
    "materials and methods": "methods",
    "experimental": "methods",
    "experimental section": "methods",
    "results": "results",
    "results and discussion": "results",
    "discussion": "discussion",
    "conclusion": "conclusion",
    "conclusions": "conclusion",
    "summary": "conclusion",
    "references": "references",
    "acknowledgements": "references",
    "acknowledgments": "references",
}


@dataclass
class Section:
    title: str
    stype: str
    # (global_paragraph_index, text)
    paras: list[tuple[int, str]] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(t for _, t in self.paras)


def _classify_heading(text: str) -> str | None:
    key = text.strip().lower().rstrip(":")
    key = key.lstrip("0123456789. ").strip()
    return SECTION_KEYWORDS.get(key)


def _looks_like_heading(p: Paragraph) -> bool:
    style = (p.style.name or "").lower() if p.style else ""
    if style.startswith("heading") or style == "title":
        return True
    text = p.text.strip()
    if 0 < len(text) <= 40 and _classify_heading(text) is not None:
        return True
    return False


def read_sections(path: Path) -> tuple[Document, list[Section]]:
    doc = Document(str(path))
    sections: list[Section] = []
    current = Section(title="(preamble)", stype="other")

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if _looks_like_heading(p):
            if current.paras:
                sections.append(current)
            stype = _classify_heading(text) or "other"
            current = Section(title=text, stype=stype)
        else:
            current.paras.append((idx, text))

    if current.paras:
        sections.append(current)
    if not sections:
        sections.append(Section(title="(document)", stype="other"))
    return doc, sections


# --------------------------------------------------------------------------- #
# Output writers
# --------------------------------------------------------------------------- #

_SEVERITY_COLOR = {
    "content": RGBColor(0xC0, 0x00, 0x00),      # red - needs human decision
    "terminology": RGBColor(0x8A, 0x2B, 0xE2),  # purple
    "structure": RGBColor(0x00, 0x66, 0xCC),    # blue
    "clarity": RGBColor(0x1F, 0x7A, 0x1F),      # green
    "language": RGBColor(0x66, 0x66, 0x66),     # grey
}


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def write_reviewed_docx(
    draft_path: Path,
    suggestions_by_index: dict[int, list[dict]],
    out_path: Path,
) -> None:
    doc = Document(str(draft_path))
    paragraphs = list(doc.paragraphs)

    # Header banner
    if paragraphs:
        banner = paragraphs[0].insert_paragraph_before()
        r = banner.add_run(
            "AI review notes inserted below relevant paragraphs. "
            "Each note = [SEVERITY] reason -> suggested rewrite. "
            "'content' notes are flags for you to verify (nothing was changed)."
        )
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    for idx in sorted(suggestions_by_index, reverse=False):
        if idx >= len(paragraphs):
            continue
        target = paragraphs[idx]
        anchor = target
        for s in suggestions_by_index[idx]:
            note = _insert_paragraph_after(anchor)
            sev = s.get("severity", "language")
            color = _SEVERITY_COLOR.get(sev, RGBColor(0x66, 0x66, 0x66))

            tag = note.add_run(f"[{sev.upper()}] ")
            tag.bold = True
            tag.font.color.rgb = color
            tag.font.highlight_color = WD_COLOR_INDEX.YELLOW

            reason = note.add_run(s.get("reason", ""))
            reason.italic = True
            reason.font.color.rgb = color

            suggestion = s.get("suggestion", "")
            if suggestion:
                arrow = note.add_run("  →  ")
                arrow.font.color.rgb = color
                sug = note.add_run(suggestion)
                sug.font.color.rgb = color
            anchor = note

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def write_accepted_docx(
    draft_path: Path,
    accepted_by_index: dict[int, list[dict]],
    out_path: Path,
) -> None:
    """Write a clean copy of the draft with only the accepted edits applied in place.

    Applies a suggestion only when it is `applicable` and its `original` is found in the
    target paragraph. A non-empty `suggestion` replaces the phrase; an empty one deletes it
    (tidying an adjacent space). Non-applicable flags are left out — no notes are inserted.
    """
    doc = Document(str(draft_path))
    paragraphs = list(doc.paragraphs)

    for idx, suggestions in accepted_by_index.items():
        if idx >= len(paragraphs):
            continue
        para = paragraphs[idx]
        text = para.text
        for s in suggestions:
            if not s.get("applicable"):
                continue
            original = (s.get("original") or "").strip()
            suggestion = (s.get("suggestion") or "").strip()
            if not original or original not in text:
                continue
            if suggestion:
                text = text.replace(original, suggestion, 1)
            else:
                # deletion: remove the phrase plus one adjacent space where present
                for variant in (original + " ", " " + original, original):
                    if variant in text:
                        text = text.replace(variant, "", 1)
                        break
        if text != para.text:
            for run in list(para.runs):
                run._element.getparent().remove(run._element)
            para.add_run(text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def write_suggestions_md(
    draft_name: str,
    section_reviews: list[dict],
    out_path: Path,
    meta: dict | None = None,
) -> None:
    lines: list[str] = []
    lines.append(f"# Review — {draft_name}")
    lines.append("")
    if meta:
        lines.append(
            f"_Backend: **{meta.get('backend')}** · "
            f"corpus passages used: {meta.get('context_used', 0)} · "
            f"sections: {len(section_reviews)}_"
        )
        lines.append("")
    total = sum(len(sr["suggestions"]) for sr in section_reviews)
    lines.append(f"**{total} suggestions** across {len(section_reviews)} sections.")
    lines.append("")
    lines.append(
        "> Legend: `content` = verify manually (nothing changed), others are "
        "language/clarity/structure/terminology edits."
    )
    lines.append("")

    for sr in section_reviews:
        lines.append(f"## {sr['title']}  \n_({sr['stype']} · {len(sr['suggestions'])} suggestions)_")
        lines.append("")
        if not sr["suggestions"]:
            lines.append("_No changes suggested._")
            lines.append("")
            continue
        for s in sr["suggestions"]:
            lines.append(f"- **[{s.get('severity','language')}]** {s.get('reason','')}")
            orig = s.get("original", "").strip()
            sug = s.get("suggestion", "").strip()
            if orig:
                lines.append(f"  - original: `{orig}`")
            if sug:
                lines.append(f"  - suggested: `{sug}`")
        lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
