"""Generate synthetic before/after pairs for data/examples/ so Horizon 1 (before/after
learning) is testable without needing the professor's real edited drafts.

Each pair models a real editing habit: trimming filler/intensifiers, softening
overreaching claims into hedged ones, swapping informal phrasing for the group's
terminology (GFET / Dirac point / label-free), and tightening run-on sentences.

Run once: `python scripts/make_sample_examples.py`
"""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "data" / "examples"


def _doc(heading_pairs: list[tuple[str, list[str]]]) -> Document:
    d = Document()
    for heading, paras in heading_pairs:
        d.add_heading(heading, level=1)
        for p in paras:
            d.add_paragraph(p)
    return d


PAIRS: dict[str, tuple[list[tuple[str, list[str]]], list[tuple[str, list[str]]]]] = {
    "biosensor_report": (
        [
            (
                "Abstract",
                [
                    "In recent years, graphene has become very popular for a lot of sensing "
                    "applications, and in this paper we made a graphene sensor that can detect "
                    "proteins and it works really well and we believe it is significant for "
                    "diagnostics because it is obviously better than other methods.",
                ],
            ),
            (
                "Introduction",
                [
                    "It is well known that detecting biomarkers is important. Graphene is a good "
                    "material. We used it to make a sensor. There are many papers about graphene "
                    "sensors, etc.",
                    "The sensitivity was very high and clearly the best, and we think this proves "
                    "our device is significant compared to previous works that were not as good.",
                ],
            ),
            (
                "Methods",
                [
                    "We grew graphene and transferred it and then we functionalized it with a "
                    "linker molecule and attached antibodies to it in order to capture the "
                    "proteins, and we measured the electrical signal.",
                ],
            ),
            (
                "Results",
                [
                    "The Dirac point moved a lot when we added the protein. The response was "
                    "very good and increased significantly with concentration. This is a really "
                    "important finding.",
                ],
            ),
            (
                "Conclusion",
                [
                    "In conclusion, we made a very good sensor that works really well and it is "
                    "obviously useful for a lot of applications.",
                ],
            ),
        ],
        [
            (
                "Abstract",
                [
                    "Graphene field-effect transistors (GFETs) have become a widely used platform "
                    "for label-free biosensing. Here, we report a GFET-based protein sensor. These "
                    "results suggest the device may be useful for diagnostics, offering advantages "
                    "over conventional methods that we describe below.",
                ],
            ),
            (
                "Introduction",
                [
                    "Rapid, label-free detection of protein biomarkers remains a central challenge "
                    "in point-of-care diagnostics. Graphene is a promising transduction material "
                    "for field-effect biosensors due to its high carrier mobility.",
                    "The sensitivity was high, and these results suggest our device is competitive "
                    "with previously reported GFET biosensors.",
                ],
            ),
            (
                "Methods",
                [
                    "Graphene was grown by CVD, transferred onto the substrate, functionalized "
                    "with a linker molecule, and conjugated with capture antibodies to bind the "
                    "target protein; the electrical signal was then recorded.",
                ],
            ),
            (
                "Results",
                [
                    "The Dirac point shifted upon protein binding. The response increased with "
                    "concentration, consistent with specific antibody-antigen binding at the "
                    "graphene surface.",
                ],
            ),
            (
                "Conclusion",
                [
                    "In conclusion, we demonstrated a GFET-based, label-free protein sensor that "
                    "may be applicable to point-of-care diagnostics.",
                ],
            ),
        ],
    ),
    "photodetector_note": (
        [
            (
                "Introduction",
                [
                    "Photodetectors are very important for a lot of applications, and obviously "
                    "graphene is a great material for them. In this paper we made a photodetector "
                    "and it works really well.",
                ],
            ),
            (
                "Results",
                [
                    "The responsivity was very high and it clearly proves that our device is "
                    "significant compared to a lot of other devices in this area.",
                    "We think the zero-bias operation is a really important advantage that makes "
                    "our device obviously better for practical use.",
                ],
            ),
            (
                "Discussion",
                [
                    "This is definitely the best result in this area and it obviously means our "
                    "approach is the right one for future work, etc.",
                ],
            ),
        ],
        [
            (
                "Introduction",
                [
                    "Photodetectors are a key building block for on-chip optical interconnects. "
                    "Graphene-based photodetectors are attractive due to their broadband "
                    "absorption and compatibility with silicon photonics. Here, we report a "
                    "waveguide-integrated graphene photodetector.",
                ],
            ),
            (
                "Results",
                [
                    "The responsivity was high, comparable to previously reported waveguide-"
                    "integrated graphene photodetectors.",
                    "The zero-bias operation is expected to reduce dark current, which may be "
                    "advantageous for low-power operation.",
                ],
            ),
            (
                "Discussion",
                [
                    "These results are consistent with efficient photocarrier generation at the "
                    "graphene-silicon Schottky junction, motivating further study of the device "
                    "under waveguide-coupled illumination.",
                ],
            ),
        ],
    ),
}


def main() -> None:
    from professor_assistant.examples import SYNTHETIC_MARKER_NAME

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for stem, (before_paras, after_paras) in PAIRS.items():
        before_path = OUT_DIR / f"{stem}.before.docx"
        after_path = OUT_DIR / f"{stem}.after.docx"
        _doc(before_paras).save(str(before_path))
        _doc(after_paras).save(str(after_path))
        print(f"Wrote {before_path}")
        print(f"Wrote {after_path}")
    # Mark the directory as synthetic so the style card / prompts never present these
    # invented pairs as the professor's real editing decisions (see examples.py).
    (OUT_DIR / SYNTHETIC_MARKER_NAME).write_text(
        "This marker means data/examples/ currently holds SYNTHETIC demo pairs generated "
        "by scripts/make_sample_examples.py, not the professor's real before/after edits.\n"
        "Delete this file once you replace them with real pairs (or it will be deleted "
        "automatically the next time this script runs, since it only ever describes the "
        "pairs this script itself wrote).\n",
        encoding="utf-8",
    )
    print(f"\n{len(PAIRS)} before/after pair(s) in {OUT_DIR}. Now run: profa examples")
    print(
        "Note: these are SYNTHETIC demo pairs for testing the pipeline end-to-end — the "
        "style card and reviewer will label anything derived from them as demo/synthetic "
        "until you replace them with the professor's real edited drafts."
    )


if __name__ == "__main__":
    main()
