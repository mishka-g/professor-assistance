"""Generate the matched professor sample paper for the golden-path demo pack.

Run: `python scripts/make_sample_paper.py`
Writes: `example_paperworks/professor_paper/professor_paper.docx`
"""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "example_paperworks" / "professor_paper" / "professor_paper.docx"

doc = Document()
doc.add_heading(
    "Graphene Field-Effect Biosensors for Label-Free Protein Detection",
    level=0,
)

doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "We report a graphene field-effect transistor (GFET) biosensor for the label-free "
    "detection of target proteins at clinically relevant concentrations. The graphene "
    "channel is functionalized with a pyrene-based linker to immobilize capture antibodies "
    "while preserving carrier mobility. Analyte binding shifts the charge-neutrality "
    "(Dirac) point, yielding a dose-dependent electrical response. The device achieves a "
    "limit of detection in the picomolar range with high selectivity against non-target "
    "proteins, demonstrating the promise of GFET platforms for point-of-care diagnostics."
)

doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "Sensitive quantification of protein biomarkers underpins early disease diagnosis. "
    "Conventional immunoassays such as ELISA are sensitive but require labeling, multiple "
    "wash steps, and bench-top instrumentation, limiting their use at the point of care. "
    "Field-effect transistors transduce molecular binding directly into an electrical "
    "signal, enabling label-free, real-time readout in a compact format."
)
doc.add_paragraph(
    "Graphene is an attractive transducer material: its two-dimensional structure places "
    "every carbon atom at the surface, so charge transport is exquisitely sensitive to "
    "interfacial perturbations. Its high carrier mobility and chemical stability further "
    "favor low-noise, reproducible operation. Prior work has established graphene as a "
    "sensitive channel, but reproducible functionalization and quantitative calibration "
    "remain the principal challenges."
)

doc.add_heading("Methods", level=1)
doc.add_paragraph(
    "Monolayer graphene was grown by chemical vapor deposition on copper foil and "
    "transferred onto a silicon dioxide substrate using a poly(methyl methacrylate) "
    "support layer. Source and drain electrodes were patterned by photolithography. The "
    "channel was functionalized with 1-pyrenebutanoic acid succinimidyl ester, which "
    "anchors to graphene through pi-pi stacking and presents a reactive ester for covalent "
    "antibody attachment."
)
doc.add_paragraph(
    "Capture antibodies specific to the target protein were immobilized on the linker, "
    "and unreacted sites were blocked to suppress non-specific adsorption. Transfer "
    "characteristics were recorded in phosphate-buffered saline while the target protein "
    "was introduced at increasing concentrations. The shift of the Dirac point was "
    "extracted as the sensing metric."
)

doc.add_heading("Results and Discussion", level=1)
doc.add_paragraph(
    "Introduction of the target protein shifted the Dirac point monotonically with "
    "concentration, consistent with charge transfer and local gating by the bound analyte. "
    "The response was well described by a Langmuir binding isotherm, from which the "
    "dissociation constant was estimated. The limit of detection reached the picomolar "
    "range, and control experiments with non-target proteins produced negligible shifts, "
    "confirming the selectivity conferred by the antibody layer."
)
doc.add_paragraph(
    "Device-to-device variation was dominated by differences in transfer quality and "
    "residual photoresist, underscoring the importance of clean processing. These results "
    "indicate that careful surface chemistry, rather than intrinsic material limits, "
    "governs sensor performance."
)

doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "We demonstrated a graphene field-effect biosensor that detects target proteins "
    "without labels at picomolar sensitivity and with high selectivity. The approach is "
    "compatible with scalable fabrication and multiplexing, and future work will address "
    "long-term stability and validation in complex biological matrices toward practical "
    "point-of-care diagnostics."
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Wrote {OUT}")
