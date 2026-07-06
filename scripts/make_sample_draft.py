"""Generate a synthetic (deliberately rough) student draft .docx for testing the pipeline.

Run once: `python scripts/make_sample_draft.py`
"""

from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "data" / "drafts" / "sample_draft.docx"

doc = Document()
doc.add_heading("A Study of Graphene Based Sensors for Detecting Proteins", level=0)

doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "In recent years, graphene has become very popular for a lot of sensing applications, "
    "and in this paper we made a graphene sensor that can detect proteins and it works "
    "really well and we believe it is significant for diagnostics because it is obviously "
    "better than other methods, and it has a lot of advantages that we will describe in "
    "order to show the performance which we measured carefully in the lab using different "
    "concentrations of the target protein over a wide range of values."
)

doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "It is well known that detecting biomarkers is important. Graphene is a good material. "
    "We used it to make a sensor. There are many papers about graphene sensors, etc."
)
doc.add_paragraph(
    "The sensitivity was very high and clearly the best, and we think this proves our "
    "device is significant compared to previous works that were not as good."
)

doc.add_heading("Methods", level=1)
doc.add_paragraph(
    "We grew graphene and transferred it and then we functionalized it with a linker "
    "molecule and attached antibodies to it in order to capture the proteins , and we "
    "measured the electrical signal."
)

doc.add_heading("Results", level=1)
doc.add_paragraph(
    "The Dirac point moved a lot when we added the protein. The response was very good and "
    "increased significantly with concentration. This is a really important finding."
)

doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "In conclusion, we made a very good sensor that works really well and it is obviously "
    "useful for a lot of applications."
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Wrote {OUT}")
